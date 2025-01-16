#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""=================================================
    @Project: FastAPIDemo
    @File： docs.py
    @Author：liaozhimingandy
    @Email: liaozhimingandy@gmail.com
    @Date：2025/1/13 10:37
    @Desc: 
================================================="""
import enum
import glob
import importlib.util
import os
from typing import List
from sqlmodel import SQLModel, Field
from enum import Enum, EnumMeta
from docx import Document

def get_model_classes_from_file(file_path: str) -> List[SQLModel]:
    """从模型文件中导入所有继承自SQLModel的类"""
    # 加载文件模块
    module_name = os.path.basename(file_path).replace(".py", "")
    spec = importlib.util.spec_from_file_location(module_name, file_path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)

    # 获取模块中的所有SQLModel子类
    model_classes = []
    for attribute_name in dir(module):
        attribute = getattr(module, attribute_name)
        if isinstance(attribute, type) and issubclass(attribute, SQLModel):
            if attribute.__name__ not in ["BaseModel", "SQLModel"]:  # 排除基础类
                model_classes.append(attribute)
    return model_classes

def get_field_info(model: SQLModel):
    """提取模型的字段信息，包括名称、类型、描述、枚举值和备注"""
    field_info = []
    for field_name, field in model.__annotations__.items():
        # 获取字段的类型和描述
        field_type = str(field)
        description = getattr(field, 'description', '')
        enum_values = []
        sa_column_kwargs = getattr(field, 'sa_column_kwargs', {})

        # 如果字段是枚举类型，获取枚举值
        if isinstance(field, type) and issubclass(field, Enum):
            enum_values = [e.value for e in field]

        field_info.append({
            "name": field_name,
            "type": field_type,
            "description": description,
            "enum_values": enum_values
        })
    return field_info


def generate_word_doc(models: list, output_file:str):
    # 创建一个Word文档对象
    doc = Document()
    doc.add_heading('FastAPI Models Documentation', 0)

    # 遍历每个模型
    for model in models:
        doc.add_heading(model.__name__, level=1)  # 模型名称作为子标题

        # 获取模型的字段信息
        fields = model.model_fields

        # 添加表格
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '序号'
        hdr_cells[1].text = '字段名称'
        hdr_cells[2].text = '字段描述'
        hdr_cells[3].text = '字段类型'
        hdr_cells[4].text = '枚举值'

        # 为每个字段添加内容
        for idx, (field_name, field_info) in enumerate(fields.items(), 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)  # 序号
            row_cells[1].text = field_name  # 字段名称
            # 获取字段描述信息
            description = field_info.description if field_info.description else ''
            row_cells[2].text = description  # 字段描述

            # 处理字段类型
            field_type = str(field_info.annotation)
            row_cells[3].text = field_type  # 字段类型

            # 枚举值
            if isinstance(field_info.annotation, enum.EnumType):
                # 检查是否是枚举类型
                enum_values = [str(e.value) for e in field_info.annotation]
                row_cells[4].text = ", ".join(enum_values)
            else:
                row_cells[4].text = ''  # 如果没有枚举值则为空

    # 保存文档
    doc.save(output_file)
    print(f"Word文档已生成：{output_file}")

if __name__ == "__main__":
    # 模型文件路径
    models_file_path = '../model/cda.py'  # 替换成实际的文件路径

    # 从模型文件中获取所有模型类
    models = get_model_classes_from_file(models_file_path)

    # 输出 Word 文件路径
    output_file = 'models_info.docx'

    # 生成 Word 文档
    generate_word_doc(models, output_file)


